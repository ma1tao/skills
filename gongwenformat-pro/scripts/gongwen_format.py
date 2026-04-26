#!/usr/bin/env python3
"""
党政机关公文标准排版脚本 (GB/T 9704-2012)
支持多种输入格式：Markdown(.md)、纯文本(.txt)、Word(.docx)、直接文本

用法:
    python3 gongwen_format.py --title "标题" --input content.md --output output.docx
    python3 gongwen_format.py --title "标题" --input content.txt --output output.docx --author "XX镇人民政府" --date "2026-04-16"
    python3 gongwen_format.py --title "标题" --input content.txt --output output.docx --print-author "XX镇人民政府办公室" --print-date "2026-04-17"
"""

import re
import argparse
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ========== 字体字号映射 ==========
FONT_TITLE = 'FZXiaoBiaoSong-B05S'
FONT_HEITI = 'SimHei'
FONT_KAITI = 'KaiTi'
FONT_FANGSONG = 'FangSong'
FONT_SONGTI = 'SimSun'
FONT_TNR = 'Times New Roman'

SIZE_ERHAO = Pt(22)
SIZE_SANHAO = Pt(16)
SIZE_SIHAO = Pt(14)

LEVEL1_PATTERNS = [re.compile(r'^[一二三四五六七八九十]+、'), re.compile(r'^第[一二三四五六七八九十]+[部分章节篇]')]
LEVEL2_PATTERNS = [re.compile(r'^[（\(][一二三四五六七八九十]+[）\)]')]
LEVEL3_PATTERNS = [re.compile(r'^\d+[\.．]')]
LEVEL4_PATTERNS = [re.compile(r'^[（\(]\d+[）\)]')]
ATTACHMENT_PATTERN = re.compile(r'^附件[：:]\s*(.*)')


def format_date(date_str):
    date_str = date_str.strip()
    if re.match(r'^\d{4}年\d{1,2}月\d{1,2}日', date_str):
        return date_str
    m = re.match(r'^(\d{4})[-/](\d{1,2})[-/](\d{1,2})', date_str)
    if m:
        return f'{m.group(1)}年{int(m.group(2))}月{int(m.group(3))}日'
    return date_str


def set_page_layout(section):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)


def _set_run_font(run, cn_font, size, bold=False):
    run.font.size = size
    run.font.bold = bold
    run.font.name = FONT_TNR
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_TNR}" w:hAnsi="{FONT_TNR}" w:eastAsia="{cn_font}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT_TNR)
        rFonts.set(qn('w:hAnsi'), cn_font)
        rFonts.set(qn('w:eastAsia'), cn_font)


def _set_songti(run):
    run.font.size = SIZE_SIHAO
    run.font.name = FONT_SONGTI
    rpr = run._element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{FONT_SONGTI}" w:hAnsi="{FONT_SONGTI}" w:eastAsia="{FONT_SONGTI}"/>')
        rpr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:ascii'), FONT_SONGTI)
        rFonts.set(qn('w:hAnsi'), FONT_SONGTI)
        rFonts.set(qn('w:eastAsia'), FONT_SONGTI)


def add_redhead(doc, org_name, doc_number=''):  
    """添加红头：发文机关标志 + 红色分隔线 + 发文字号"""
    # 上边缘至版心上边缘为35mm，版心上边缘=上边距37mm，所以红头顶边距=37-35=2mm
    # 通过段前间距控制：35mm ≈ 99pt（上边距已留37mm，再留35mm-字号高度的空间）
    # 实际做法：让红头段落从页边距下方35mm-37mm处开始，用space_before控制

    # 发文机关标志（红色小标宋，居中）
    display_name = org_name
    p_org = doc.add_paragraph()
    p_org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_org.paragraph_format.space_before = Pt(0)
    p_org.paragraph_format.space_after = Pt(0)
    p_org.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_org.paragraph_format.line_spacing = Pt(40)
    run_org = p_org.add_run(display_name)
    # 字号根据名称字数自适应
    name_len = len(display_name)
    if name_len <= 6:
        font_size = Pt(32)
    elif name_len <= 10:
        font_size = Pt(28)
    elif name_len <= 15:
        font_size = Pt(24)
    else:
        font_size = Pt(22)
    _set_run_font(run_org, FONT_TITLE, font_size, bold=True)  # 方正小标宋简体
    run_org.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色

    # 发文字号（三号仿宋，居中）
    if doc_number:
        p_num = doc.add_paragraph()
        p_num.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_num.paragraph_format.space_before = Pt(0)
        p_num.paragraph_format.space_after = Pt(0)
        p_num.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p_num.paragraph_format.line_spacing = Pt(28)
        run_num = p_num.add_run(doc_number)
        _set_run_font(run_num, FONT_FANGSONG, SIZE_SANHAO)

    # 红色分隔线（156mm宽）
    p_line = doc.add_paragraph()
    p_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_line.paragraph_format.space_before = Pt(0)
    p_line.paragraph_format.space_after = Pt(0)
    p_line.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p_line.paragraph_format.line_spacing = Pt(4)
    pPr = p_line._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="24" w:space="1" w:color="FF0000"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)
    p_line.paragraph_format.left_indent = Cm(0)
    p_line.paragraph_format.right_indent = Cm(0)

    # 分隔线与标题之间空一行
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Pt(28)


def add_title(doc, title_text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(32)
    p.paragraph_format.first_line_indent = None
    run = p.add_run(title_text)
    _set_run_font(run, FONT_TITLE, SIZE_ERHAO, bold=True)
    return p


def add_body_paragraph(doc, text, level=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(28)
    p.paragraph_format.first_line_indent = Cm(1.13)
    font_map = {0: FONT_FANGSONG, 1: FONT_HEITI, 2: FONT_KAITI, 3: FONT_FANGSONG, 4: FONT_FANGSONG}
    run = p.add_run(text)
    _set_run_font(run, font_map.get(level, FONT_FANGSONG), SIZE_SANHAO)
    return p


def add_attachment_block(doc, lines):
    """添加附件说明块（GB/T 9704-2012 §7.3.4）。
    规范要求：
    - 正文下空一行
    - 左空二字编排（首行缩进2字符）
    - 多个附件每行一个，每个附件名称后不加标点
    - 续行（第2个附件起）与第一个附件名称左对齐
    """
    if not lines:
        return
    # 正文下空一行
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Pt(28)

    for i, line in enumerate(lines):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(28)
        # 国标：左空二字 = 首行缩进2字符
        if i == 0:
            # 第一行「附件：1. xxx」左空二字
            p.paragraph_format.first_line_indent = Cm(1.13)
        else:
            # 续行：与「附件」后的内容对齐，即左空二字+「附件：」宽度
            # 「附件：」= 2个中文字+1个全角冒号 = 3字符 ≈ Cm(1.13) + Cm(0.85)
            # 简化：直接左空二字，前面手动补缩进用空格
            p.paragraph_format.first_line_indent = Cm(1.13)
            # 在行首加空格对齐到「附件：」之后
            line = '　　' + line  # 用2个全角空格模拟对齐
        run = p.add_run(line)
        _set_run_font(run, FONT_FANGSONG, SIZE_SANHAO)


def _parse_attachment_content(text):
    """解析附件冒号后的内容，智能拆分为多行。
    '1. xxx  2. xxx' -> ['附件：1. xxx', '2. xxx']
    '关于XXX的通知'   -> ['附件：关于XXX的通知']
    注意：附件名称后不加标点符号（国标要求）
    """
    text = text.strip()
    # 去除附件名称末尾的标点
    text = re.sub(r'[；;，,。.！!？?]+$', '', text)
    if not text:
        return ['附件：（见附件）']
    # 检测多个编号项
    items = re.findall(r'(\d+[\.．][\s\S]*?)(?=\s+\d+[\.．]|$)', text)
    if items and len(items) > 1:
        lines = []
        for j, item in enumerate(items):
            item = item.strip()
            item = re.sub(r'[；;，,。.！!？?]+$', '', item)  # 去末尾标点
            if j == 0:
                lines.append(f'附件：{item}')
            else:
                lines.append(item)
        return lines
    return [f'附件：{text}']


def _parse_attachment_lines(lines, start_idx):
    """从 start_idx 解析附件块，返回 (attachment_lines, next_idx)。
    识别续行：缩进行、数字编号行、附件N：行。
    """
    result = [lines[start_idx]]
    idx = start_idx + 1
    while idx < len(lines):
        stripped = lines[idx].strip()
        if not stripped:
            idx += 1
            continue
        is_cont = (
            lines[idx].startswith('  ') or
            lines[idx].startswith('\t') or
            re.match(r'^\s*\d+[\.．〔〕]', stripped) or
            re.match(r'^附件\d+[：:]', stripped)
        )
        if is_cont:
            cleaned = re.sub(r'^\s{1,2}', '', stripped)
            result.append(cleaned)
            idx += 1
        else:
            break
    return result, idx


def _build_footer_xml(alignment):
    """构建页脚段落XML：—页码—，四号宋体"""
    return (
        f'<w:p {nsdecls("w")}>'
        f'<w:pPr><w:jc w:val="{alignment}"/></w:pPr>'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT_SONGTI}" w:hAnsi="{FONT_SONGTI}" w:eastAsia="{FONT_SONGTI}"/>'
        f'<w:sz w:val="28"/><w:szCs w:val="28"/>'
        f'</w:rPr><w:t>\u2014</w:t></w:r>'
        f'<w:fldSimple w:instr=" PAGE "><w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT_SONGTI}" w:hAnsi="{FONT_SONGTI}" w:eastAsia="{FONT_SONGTI}"/>'
        f'<w:sz w:val="28"/><w:szCs w:val="28"/>'
        f'</w:rPr><w:t>1</w:t></w:r></w:fldSimple>'
        f'<w:r><w:rPr>'
        f'<w:rFonts w:ascii="{FONT_SONGTI}" w:hAnsi="{FONT_SONGTI}" w:eastAsia="{FONT_SONGTI}"/>'
        f'<w:sz w:val="28"/><w:szCs w:val="28"/>'
        f'</w:rPr><w:t>\u2014</w:t></w:r>'
        f'</w:p>'
    )


def add_page_number(doc, skip_first=False):
    """添加页码：四号宋体，—1— 格式，单页居右、双页居左。
    skip_first=True 时启用首页不同，首页不显示页码（红头文件标准）。"""
    # 启用奇偶页不同
    docSettings = doc.settings.element
    if docSettings.find(qn('w:evenAndOddHeaders')) is None:
        docSettings.append(parse_xml(f'<w:evenAndOddHeaders {nsdecls("w")}/>'))
    # 启用首页不同（红头文件首页不显示页码）
    if skip_first:
        sectPr = doc.sections[0]._sectPr
        if sectPr.find(qn('w:titlePg')) is None:
            sectPr.append(parse_xml(f'<w:titlePg {nsdecls("w")}/>'))

    # 清理所有 section 的 footer
    for section in doc.sections:
        sectPr = section._sectPr
        for ref in list(sectPr.findall(qn('w:footerReference'))):
            sectPr.remove(ref)
        for tag in ['w:oddFooter', 'w:evenFooter']:
            for el in list(sectPr.findall(qn(tag))):
                sectPr.remove(el)
        section.footer.is_linked_to_previous = True

    # 第一个 section：奇数页页码（居右）
    s = doc.sections[0]
    s.footer.is_linked_to_previous = False
    p_odd = s.footer.paragraphs[0]
    p_odd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in list(p_odd.runs):
        p_odd._element.remove(run._element)
    r1 = p_odd.add_run('\u2014')
    r1.font.name = FONT_SONGTI
    r1.font.size = SIZE_SIHAO
    r2 = p_odd.add_run()
    r2.font.name = FONT_SONGTI
    r2.font.size = SIZE_SIHAO
    r2._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r2._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r2._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    r3 = p_odd.add_run('\u2014')
    r3.font.name = FONT_SONGTI
    r3.font.size = SIZE_SIHAO

    # 第一个 section：偶数页页码（居左）
    even_footer = s.even_page_footer
    even_footer.is_linked_to_previous = False
    p_even = even_footer.paragraphs[0]
    p_even.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r4 = p_even.add_run('\u2014')
    r4.font.name = FONT_SONGTI
    r4.font.size = SIZE_SIHAO
    r5 = p_even.add_run()
    r5.font.name = FONT_SONGTI
    r5.font.size = SIZE_SIHAO
    r5._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r5._element.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
    r5._element.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    r6 = p_even.add_run('\u2014')
    r6.font.name = FONT_SONGTI
    r6.font.size = SIZE_SIHAO

    # 版记 section（最后一个）不设页码，保持 linked_to_previous=False 但内容清空
    if len(doc.sections) > 1:
        last = doc.sections[-1]
        last.footer.is_linked_to_previous = False
        last.even_page_footer.is_linked_to_previous = False
        for p in last.footer.paragraphs:
            for run in list(p.runs):
                p._element.remove(run._element)
        for p in last.even_page_footer.paragraphs:
            for run in list(p.runs):
                p._element.remove(run._element)

# ========== 输入格式处理 ==========

def detect_level_from_text(text):
    text = text.strip()
    if not text:
        return None
    for pat in LEVEL1_PATTERNS:
        if pat.match(text): return 1
    for pat in LEVEL2_PATTERNS:
        if pat.match(text): return 2
    for pat in LEVEL3_PATTERNS:
        if pat.match(text): return 3
    for pat in LEVEL4_PATTERNS:
        if pat.match(text): return 4
    return 0


def detect_level(line):
    line_stripped = line.strip()
    if not line_stripped:
        return None, None
    if line_stripped.startswith('## ') and not line_stripped.startswith('### '):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', line_stripped[3:].strip())
        return 1, text
    if line_stripped.startswith('### '):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', line_stripped[4:].strip())
        return 2, text
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line_stripped).strip()
    if not clean:
        return None, None
    level = detect_level_from_text(clean)
    return level, clean


def normalize_content(content):
    result, i = [], 0
    while i < len(content):
        if content[i] == '"' and i + 1 < len(content):
            close = content.find('"', i + 1)
            if close != -1:
                result.append('\u201c')
                result.append(content[i+1:close])
                result.append('\u201d')
                i = close + 1
                continue
        result.append(content[i])
        i += 1
    return ''.join(result)


def read_input(filepath):
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        doc = Document(filepath)
        lines = [para.text.strip() for para in doc.paragraphs if para.text.strip()]
        return normalize_content('\n'.join(lines)), False
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    return normalize_content(content), (ext == '.md')


def split_heading_and_body(text, level):
    if level != 2:
        return [(text, level)]
    dot_pos = text.find('。')
    if dot_pos == -1:
        return [(text, level)]
    heading, body = text[:dot_pos + 1], text[dot_pos + 1:].strip()
    if not body:
        return [(text, level)]
    return [(heading, level), (body, 0)]


def parse_and_add_content(doc, content):
    lines = content.split('\n')
    i = 0
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        att_match = ATTACHMENT_PATTERN.match(stripped)
        if att_match:
            att_lines, next_i = _parse_attachment_lines(lines, i)
            first_content = ATTACHMENT_PATTERN.match(att_lines[0]).group(1).strip()
            parsed = _parse_attachment_content(first_content)
            if len(att_lines) > 1:
                for extra in att_lines[1:]:
                    parsed.append(extra.strip())
            add_attachment_block(doc, parsed)
            i = next_i
            continue
        level, text = detect_level(lines[i])
        if level is None or not text:
            i += 1
            continue
        for part_text, part_level in split_heading_and_body(text, level):
            add_body_paragraph(doc, part_text, level=part_level)
        i += 1


def main():
    parser = argparse.ArgumentParser(description='党政机关公文标准排版')
    parser.add_argument('--title', required=True, help='公文标题')
    parser.add_argument('--input', required=True, help='输入文件（.md / .txt / .docx）')
    parser.add_argument('--output', required=True, help='输出Word文件路径')
    parser.add_argument('--author', default='', help='发文机关名称')
    parser.add_argument('--date', default='', help='成文日期（自动格式化）')
    parser.add_argument('--print-author', default='', help='印发机关（版记）')
    parser.add_argument('--print-date', default='', help='印发日期（自动格式化）')
    parser.add_argument('--cc', default='', help='抄送机关（版记中）')
    parser.add_argument('--redhead', default='', help='红头机关名称（如 XX省人民政府）')
    parser.add_argument('--doc-number', default='', help='发文字号（如 X政发〔2026〕X号）')
    parser.add_argument('--no-page-num', action='store_true', help='不添加页码（默认添加，红头文件首页自动跳过）')
    args = parser.parse_args()

    content, _ = read_input(args.input)

    doc = Document()
    section = doc.sections[0]
    set_page_layout(section)

    # 红头（在标题之前）
    if args.redhead:
        add_redhead(doc, args.redhead, args.doc_number)

    add_title(doc, args.title)

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(0)
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    spacer.paragraph_format.line_spacing = Pt(28)

    parse_and_add_content(doc, content)

    if args.author or args.date:
        for _ in range(3):
            doc.add_paragraph()
        if args.author:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_indent = Cm(1.3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(28)
            run = p.add_run(args.author)
            _set_run_font(run, FONT_FANGSONG, SIZE_SANHAO)
        if args.date:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.right_indent = Cm(1.3)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(28)
            run = p.add_run(format_date(args.date))
            _set_run_font(run, FONT_FANGSONG, SIZE_SANHAO)

    # 版记（新建分节，垂直对齐到底部，上下各一条反线）
    if args.print_author or args.print_date:
        from docx.enum.section import WD_SECTION_START
        new_section = doc.add_section(WD_SECTION_START.NEW_PAGE)
        set_page_layout(new_section)
        # 垂直对齐到底部
        new_section._sectPr.append(parse_xml(f'<w:vAlign {nsdecls("w")} w:val="bottom"/>'))

        def _add_banji_para(text, alignment=WD_ALIGN_PARAGRAPH.LEFT, has_top_border=False, has_bottom_border=False):
            p = doc.add_paragraph()
            p.alignment = alignment
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(28)
            if text:
                run = p.add_run(text)
                _set_run_font(run, FONT_FANGSONG, SIZE_SANHAO)
            if has_top_border or has_bottom_border:
                pPr = p._element.get_or_add_pPr()
                borders = []
                if has_top_border:
                    borders.append(f'<w:top w:val="single" w:sz="4" w:space="1" w:color="000000"/>')
                if has_bottom_border:
                    borders.append(f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="000000"/>')
                if borders:
                    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>{"".join(borders)}</w:pBdr>')
                    pPr.append(pBdr)
            return p

        # 上反线
        _add_banji_para('', has_bottom_border=True)
        # 抄送机关
        cc = getattr(args, 'cc', '')
        if cc:
            _add_banji_para(f'抄送：{cc}')
        # 印发机关 + 印发日期
        print_text = ''
        if args.print_author and args.print_date:
            print_text = f'{args.print_author}            {format_date(args.print_date)}印发'
        elif args.print_author:
            print_text = args.print_author
        elif args.print_date:
            print_text = f'{format_date(args.print_date)}印发'
        if print_text:
            _add_banji_para(print_text, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        # 下反线
        _add_banji_para('', has_top_border=True)

    if not args.no_page_num:
        add_page_number(doc, skip_first=bool(args.redhead))

    out_dir = os.path.dirname(os.path.abspath(args.output))
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(args.output)
    print(f"✅ 公文已生成: {args.output}")


if __name__ == '__main__':
    main()
